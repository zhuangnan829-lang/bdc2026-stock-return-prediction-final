from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT_DIR = Path(__file__).resolve().parents[3]
OUT_DIR = ROOT_DIR / "app" / "model" / "final_process_summary"
DOCX_PATH = OUT_DIR / "收益率与稳定性优化全过程总结报告.docx"
MD_PATH = OUT_DIR / "收益率与稳定性优化全过程总结报告.md"


TITLE = "收益率与稳定性优化全过程总结报告"
SUBTITLE = "基于 Codex 分步执行提示词清单与项目最终封板产物"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
        set_cell_shading(table.rows[0].cells[idx], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    document.add_paragraph()


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def configure_styles(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(31, 78, 121)),
        ("Heading 2", 13, RGBColor(47, 84, 150)),
        ("Heading 3", 11, RGBColor(31, 78, 121)),
    ]:
        style = document.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True


def add_cover(document: Document) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(31, 78, 121)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(SUBTITLE)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(89, 89, 89)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    document.add_paragraph()
    add_table(
        document,
        ["项目项", "最终状态"],
        [
            ["主线模型", "LSTM sl20"],
            ["最终增强", "HV rerank，高波动阶段 close_position_20d 权重 -0.05"],
            ["当前默认 profile", "lstm_sl20_base_alpha_v3_rs_crowding_mini4__hv_close_position_rerank"],
            ["最终 result.csv", "300316、600115、600183、600584、688396，权重均为 0.18"],
            ["最终提交包", "THU-BDC2026-hv-rerank-final_20260525_233344.zip"],
            ["验证状态", "完整验证、validator、pre-submit、解包复跑均通过"],
        ],
        widths=[4.2, 11.5],
    )
    document.add_section(WD_SECTION_START.NEW_PAGE)


def process_rows() -> list[list[str]]:
    return [
        ["Prompt 0", "实验目录规范与总控", "新增实验目录标准、experiment_utils、排行榜字段规范。", "让实验可登记、可追踪、可复盘，避免结果散落。"],
        ["Prompt 1", "统一配置来源", "建立 default_submission_config 权威源、load_submission_config、配置一致性检查。", "消除脚本、JSON、快照之间的配置漂移。"],
        ["Prompt 2", "收益瓶颈拆解", "拆解候选池、排序、权重、单票贡献和 true top missed 情况。", "先判断问题来自哪里，避免盲目换模型。"],
        ["Prompt 3", "权重上限搜索", "搜索 max_single_weight，比较收益、回撤、换手和集中度。", "证明 cap 能降风险，但收益会受影响。"],
        ["Prompt 4", "pred/equal 混合权重", "测试 pred、equal、blend_0.5 等权重组合。", "给 robust 配置提供更稳的权重候选。"],
        ["Prompt 5", "换手压力测试", "测试 max_turnover、交易成本、权重组合压力。", "明确低换手方案能降回撤，但收益弹性下降。"],
        ["Prompt 6", "RankIC 稳定性", "评估 rank_ic_mean、worst fold、fold 间稳定性。", "不再只看平均值，开始关注最差折风险。"],
        ["Prompt 7", "误排样本诊断", "定位高预测低真实、低预测高真实样本特征。", "为 rerank、风险过滤和特征修正提供依据。"],
        ["Prompt 8", "市场状态评估", "按低波动、高波动、趋势、震荡、高波动震荡拆分表现。", "发现高波动震荡阶段是明显弱点。"],
        ["Prompt 9", "Regime aware 思路", "评估不同市场状态下是否需要 aggressive/robust 分流。", "为后续 regime rerank 和 robust 配置打基础。"],
        ["Prompt 10", "特征集同协议对比", "对 base_alpha_v4_medium 等特征集做同协议比较。", "避免只看训练效果，强调统一回测口径。"],
        ["Prompt 11", "Top-K 加权 Rank 目标", "新增 Top-K 加权训练目标实验。", "尝试让训练目标更贴近最终 Top5 收益。"],
        ["Prompt 12", "标签变体实验", "测试 clipped、residual、risk_adjusted 等标签。", "丰富候选，但用硬规则防止过拟合采用。"],
        ["Prompt 13", "多随机种子 Bagging", "实现多 seed 稳定性候选。", "降低单模型随机性，但不直接替换主线。"],
        ["Prompt 14", "Snapshot Ensemble", "测试训练快照平均。", "评估 epoch 偶然性对稳定性的影响。"],
        ["Prompt 15", "Rank Blend", "实现 LSTM、LightGBM、Momentum、XGBoost 等 rank 融合。", "验证融合能否改善稳定性或收益，结果未足以替换主线。"],
        ["Prompt 16", "sl40/sl60 长序列", "建立长序列候选分支。", "借鉴长窗口思路，但按规则只作为候选。"],
        ["Prompt 17", "Transformer-lite", "实现轻量 Transformer 分支。", "探索模型替代可能，但未全面胜出，未替换 sl20。"],
        ["Prompt 18", "Regime Switching", "建立 aggressive/robust 状态切换实验。", "证明全配置切换过重，容易牺牲收益。"],
        ["Prompt 19", "aggressive/robust 双配置与选择器", "新增 submission_aggressive、submission_robust、最终选择报告。", "形成冲分/稳健两套候选，不直接覆盖默认。"],
        ["Prompt 20", "实验排行榜和采用规则", "建立 stable_alpha_score、硬性淘汰规则、Top10 候选报告。", "模型选择从凭感觉变为证据驱动。"],
        ["Prompt 21", "完整验证流水线", "串联 14 个关键步骤，失败记录但继续执行。", "形成端到端验证闭环，封板时失败步骤为 0。"],
        ["Prompt 22", "高波动震荡专项优化", "运行 HVR 优化和阈值搜索，发现全配置 robust switch 不合适。", "把问题缩小到轻量 rerank，而不是整体切换。"],
        ["Prompt 23", "HV rerank 提交候选", "生成 result_hv_rerank.csv，并与当前 result 做提交级对比。", "确认只换 1 只票，且 validator/pre-submit 通过。"],
        ["Prompt 24", "最终候选决策表", "比较 current、aggressive、robust、HV rerank 四类候选。", "明确默认保留主线，HV rerank 是最值得确认的增强候选。"],
        ["人工确认", "600115 换入、601877 换出", "分析两只票的预测分、位置、波动、换手与高波动适配性。", "确认切换有利于降低高位拥挤风险。"],
        ["Prompt 25", "同步 HV rerank 默认", "备份旧配置和 result，同步 default、best_config、model_meta、submission_artifacts。", "让正式推理入口真正使用 HV rerank。"],
        ["Prompt 26", "最终封板与打包", "freeze、完整验证、最小提交包、解包复跑、SHA256 记录。", "形成可提交、可复跑、可追溯的最终交付包。"],
        ["README 完善", "交付说明文档", "更新 README 和 app/readme，写明最终方案、命令、结果和验证。", "降低复现和答辩说明成本。"],
    ]


def write_markdown() -> None:
    rows = process_rows()
    lines = [
        f"# {TITLE}",
        "",
        "## 一、总体结论",
        "",
        "本轮优化按照《Codex 分步执行提示词清单》的原则推进：不盲目推翻 LSTM sl20 主线，而是先诊断、再做组合层优化、稳定性评估、候选模型对比、最终选择器和完整验证。最终没有采用新模型替代 sl20，而是在 sl20 主线上加入高波动阶段的 HV rerank 轻量增强，并完成封板打包。",
        "",
        "最终默认方案：`lstm_sl20_base_alpha_v3_rs_crowding_mini4__hv_close_position_rerank`。",
        "",
        "最终 result.csv：`300316, 600115, 600183, 600584, 688396`，每只权重 `0.18`。",
        "",
        "最终提交包：`THU-BDC2026-hv-rerank-final_20260525_233344.zip`。",
        "",
        "## 二、对程序的核心提升",
        "",
        "- 工程可靠性提升：统一配置源，增加一致性检查、pre-submit 检查和完整验证流水线。",
        "- 模型选择更稳健：用 leaderboard、stable_alpha_score 和硬性淘汰规则替代凭感觉选模型。",
        "- 风险控制更清晰：系统评估权重上限、混合权重、低换手压力和单票集中。",
        "- 市场状态适配增强：发现高波动震荡弱点，并用 HV rerank 做轻量修正。",
        "- 提交流程闭环：最终配置、result、submission_artifacts、封板报告和提交包全部同步。",
        "",
        "## 三、关键指标变化和最终证据",
        "",
        "- 当前主线回测：cost_after_return `1.171246`，Sharpe `4.019488`，max_drawdown `-0.090067`，avg_turnover `0.956671`。",
        "- HV rerank 相对 baseline：delta_return `+0.106056`，delta_sharpe `+0.187586`，delta_max_drawdown `+0.002645`，delta_high_vol `+0.002761`。",
        "- robust 低换手候选：`mt050_tc0010_blend_0.5_cap0.20`，max_drawdown `-0.048990`，avg_turnover `0.500000`，但 return `0.653410`，因此只作为稳健观察，不作为默认。",
        "- 完整验证流水线：14 步全部 PASS，失败步骤 `0`。",
        "- 解包复跑：最终提交包解包后 `cli.py predict` 可重新生成同样的 result.csv。",
        "",
        "## 四、每一步做了什么、有什么作用",
        "",
        "| 步骤 | 主题 | 做了什么 | 作用 |",
        "|---|---|---|---|",
    ]
    for row in rows:
        lines.append("| " + " | ".join(row) + " |")
    lines.extend(
        [
            "",
            "## 五、最终采用与放弃逻辑",
            "",
            "- 采用 HV rerank：因为它不替换 sl20，只在高波动阶段轻量修正排序，且提交级验证通过。",
            "- 保留 aggressive：用于比赛冲分思路，但不直接作为默认提交。",
            "- 保留 robust：用于稳定/低换手场景，但收益下降明显，不作为本轮默认。",
            "- 放弃直接替换 sl20：因为没有新模型在 walk-forward、回测、单切片三方面同时稳定胜出。",
            "",
            "## 六、最终交付物",
            "",
            "- `app/model/default_submission_config.json`：当前权威默认配置。",
            "- `app/output/result.csv`：最终提交结果。",
            "- `app/model/full_optimization_validation_report.md`：完整验证报告。",
            "- `app/model/final_candidate_decision/final_candidate_decision_report.md`：最终候选决策表。",
            "- `app/model/final_submission_package/final_package_freeze_report.md`：封板报告。",
            "- `THU-BDC2026-hv-rerank-final_20260525_233344.zip`：最终提交包。",
            "",
            "## 七、后续建议",
            "",
            "当前已经进入封板状态，提交前不建议继续修改模型、配置或 result.csv。后续如果继续研究，建议单独开新实验分支，重点探索高波动震荡阶段的特征、标签和轻量 rerank，而不要直接覆盖当前封板版本。",
            "",
        ]
    )
    MD_PATH.write_text("\n".join(lines), encoding="utf-8")


def build_docx() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    document = Document()
    configure_styles(document)
    add_cover(document)

    document.add_heading("一、总体结论", level=1)
    document.add_paragraph(
        "本轮优化严格围绕提示词文档中的原则推进：不盲目推翻 LSTM sl20 主线，而是先诊断、再做组合层优化、"
        "稳定性评估、候选模型对比、最终选择器和完整验证。最终没有采用新模型替代 sl20，而是在 sl20 主线上加入"
        "高波动阶段的 HV rerank 轻量增强，并完成封板打包。"
    )
    add_table(
        document,
        ["项目", "最终结果"],
        [
            ["最终默认方案", "lstm_sl20_base_alpha_v3_rs_crowding_mini4__hv_close_position_rerank"],
            ["最终持仓", "300316、600115、600183、600584、688396，权重均为 0.18"],
            ["核心变化", "换入 600115，换出 601877"],
            ["最终提交包", "THU-BDC2026-hv-rerank-final_20260525_233344.zip"],
            ["验证状态", "完整验证、validator、pre-submit、解包复跑均通过"],
        ],
        widths=[4.0, 11.8],
    )

    document.add_heading("二、对程序的核心提升", level=1)
    add_bullets(
        document,
        [
            "工程可靠性提升：统一配置源，增加一致性检查、pre-submit 检查和完整验证流水线。",
            "模型选择更稳健：用 leaderboard、stable_alpha_score 和硬性淘汰规则替代凭感觉选模型。",
            "风险控制更清晰：系统评估权重上限、混合权重、低换手压力和单票集中。",
            "市场状态适配增强：发现高波动震荡弱点，并用 HV rerank 做轻量修正。",
            "提交流程闭环：最终配置、result、submission_artifacts、封板报告和提交包全部同步。",
        ],
    )

    document.add_heading("三、关键指标和最终证据", level=1)
    add_table(
        document,
        ["证据项", "结果", "含义"],
        [
            ["当前主线回测", "return 1.171246 / Sharpe 4.019488 / MDD -0.090067", "主线收益能力仍然较强，是保留 sl20 的依据。"],
            ["HV rerank 增量", "delta_return +0.106056 / delta_sharpe +0.187586", "高波动轻重排对收益和风险调整收益都有正贡献。"],
            ["高波动 Top5", "delta_high_vol +0.002761", "针对原先较弱的高波动阶段有改善。"],
            ["robust 低换手", "turnover 0.500000 / MDD -0.048990 / return 0.653410", "风险更稳但收益下降，适合作为备用而非默认。"],
            ["完整验证", "14 步 PASS，失败步骤 0", "当前提交链路可复现、可检查。"],
        ],
        widths=[3.8, 5.2, 6.6],
    )

    document.add_heading("四、每一步做了什么、有什么作用", level=1)
    add_table(document, ["步骤", "主题", "做了什么", "作用"], process_rows(), widths=[2.4, 3.3, 5.2, 5.2])

    document.add_heading("五、最终采用与放弃逻辑", level=1)
    add_bullets(
        document,
        [
            "采用 HV rerank：它不替换 sl20，只在高波动阶段轻量修正排序，且提交级验证通过。",
            "保留 aggressive：用于比赛冲分思路，但当前不直接作为默认提交。",
            "保留 robust：用于稳定/低换手场景，但收益下降明显，因此不作为本轮默认。",
            "放弃直接替换 sl20：因为没有新模型在 walk-forward、回测、单切片三方面同时稳定胜出。",
            "放弃全配置 regime switch：整体切换过重，收益损失不划算；最终采用轻量 rerank。",
        ],
    )

    document.add_heading("六、最终交付物", level=1)
    add_table(
        document,
        ["文件", "作用"],
        [
            ["app/model/default_submission_config.json", "当前权威默认配置。"],
            ["app/output/result.csv", "最终提交结果。"],
            ["app/model/full_optimization_validation_report.md", "完整验证报告。"],
            ["app/model/final_candidate_decision/final_candidate_decision_report.md", "最终候选决策表。"],
            ["app/model/final_submission_package/final_package_freeze_report.md", "封板报告。"],
            ["THU-BDC2026-hv-rerank-final_20260525_233344.zip", "最终提交包。"],
        ],
        widths=[7.0, 8.4],
    )

    document.add_heading("七、后续建议", level=1)
    document.add_paragraph(
        "当前已经进入封板状态，提交前不建议继续修改模型、配置或 result.csv。后续如果继续研究，建议单独开新实验分支，"
        "重点探索高波动震荡阶段的特征、标签和轻量 rerank，而不要直接覆盖当前封板版本。"
    )

    document.save(DOCX_PATH)


def main() -> None:
    build_docx()
    write_markdown()
    print(f"wrote {DOCX_PATH}")
    print(f"wrote {MD_PATH}")


if __name__ == "__main__":
    main()
